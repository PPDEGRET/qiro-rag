"""Document parsing for evidence packs.

Docling is attempted when installed. Lightweight local parsers keep the core CLI
usable without a heavyweight setup. OCR is optional and local-only via
pdf2image/pytesseract when installed.
"""

from __future__ import annotations

import csv
import os
import re
import shutil
from pathlib import Path
from typing import Any

from qiro_rag.schemas import ParsedChunk, ParsedDocument, ParsedTable
from qiro_rag.utils import normalize_ws

MAX_CHARS = 1400
OVERLAP = 160
IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".tif", ".tiff"}


class ParserUnavailable(RuntimeError):
    """Raised when a file cannot be parsed with installed local parsers."""


def parse_file(path: Path, doc_id: str, stored_path: str, parser: str = "auto") -> ParsedDocument:
    suffix = path.suffix.lower()
    warnings: list[str] = []

    if parser in {"auto", "docling"} and suffix in {
        ".pdf",
        ".docx",
        ".xlsx",
        ".csv",
        ".md",
        ".txt",
    }:
        parsed = _parse_with_docling(path, doc_id, stored_path)
        if parsed is not None:
            return parsed
        if parser == "docling":
            warnings.append(
                "Docling was requested but unavailable or failed; used lightweight parser."
            )

    if suffix in {".md", ".txt"}:
        text = path.read_text(encoding="utf-8", errors="replace")
        return _document_from_sections(
            stored_path, doc_id, "text", [(None, None, path.stem, text)], warnings
        )
    if suffix == ".csv":
        return _parse_csv(path, doc_id, stored_path, warnings)
    if suffix == ".xlsx":
        return _parse_xlsx(path, doc_id, stored_path, warnings)
    if suffix == ".docx":
        return _parse_docx(path, doc_id, stored_path, warnings)
    if suffix == ".pdf":
        return _parse_pdf(path, doc_id, stored_path, warnings, ocr=parser in {"auto", "ocr"})
    if suffix in IMAGE_SUFFIXES:
        return _parse_image_ocr(path, doc_id, stored_path, warnings)
    raise ParserUnavailable(f"Unsupported file type: {path.suffix}")


def _parse_with_docling(path: Path, doc_id: str, stored_path: str) -> ParsedDocument | None:
    try:
        from docling.document_converter import DocumentConverter  # type: ignore[import-not-found]
    except Exception:
        return None

    try:
        result = DocumentConverter().convert(str(path))
        document = result.document
        if hasattr(document, "export_to_markdown"):
            text = document.export_to_markdown()
        elif hasattr(document, "export_to_text"):
            text = document.export_to_text()
        else:
            text = str(document)
    except Exception:
        return None

    return _document_from_sections(
        stored_path, doc_id, "docling", [(None, None, path.stem, text)], []
    )


def _parse_csv(path: Path, doc_id: str, stored_path: str, warnings: list[str]) -> ParsedDocument:
    rows: list[str] = []
    with path.open("r", encoding="utf-8", errors="replace", newline="") as handle:
        reader = csv.reader(handle)
        for row in reader:
            rows.append(" | ".join(str(cell) for cell in row))
    text = "\n".join(rows)
    doc = _document_from_sections(
        stored_path, doc_id, "csv", [(None, path.stem, path.stem, text)], warnings
    )
    doc.tables.append(
        ParsedTable(
            table_id=f"{doc_id}-table-0001",
            doc_id=doc_id,
            path=stored_path,
            sheet=path.stem,
            text=text,
        )
    )
    return doc


def _parse_xlsx(path: Path, doc_id: str, stored_path: str, warnings: list[str]) -> ParsedDocument:
    try:
        from openpyxl import load_workbook
    except Exception as exc:  # pragma: no cover - dependency declared
        raise ParserUnavailable("openpyxl is required for .xlsx files") from exc

    workbook = load_workbook(path, read_only=True, data_only=True)
    sections: list[tuple[int | None, str | None, str | None, str]] = []
    tables: list[ParsedTable] = []
    for sheet in workbook.worksheets:
        rows: list[str] = []
        for row in sheet.iter_rows(values_only=True):
            values = ["" if cell is None else str(cell) for cell in row]
            if any(value.strip() for value in values):
                rows.append(" | ".join(values))
        text = "\n".join(rows)
        if text.strip():
            sections.append((None, sheet.title, sheet.title, text))
            tables.append(
                ParsedTable(
                    table_id=f"{doc_id}-table-{len(tables) + 1:04d}",
                    doc_id=doc_id,
                    path=stored_path,
                    sheet=sheet.title,
                    text=text,
                )
            )
    doc = _document_from_sections(stored_path, doc_id, "xlsx", sections, warnings)
    doc.tables.extend(tables)
    return doc


def _parse_docx(path: Path, doc_id: str, stored_path: str, warnings: list[str]) -> ParsedDocument:
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover - dependency declared
        raise ParserUnavailable("python-docx is required for .docx files") from exc

    document = Document(path)
    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    text = "\n".join(parts)
    return _document_from_sections(
        stored_path, doc_id, "docx", [(None, None, path.stem, text)], warnings
    )


def _parse_pdf(
    path: Path, doc_id: str, stored_path: str, warnings: list[str], ocr: bool
) -> ParsedDocument:
    try:
        from pypdf import PdfReader
    except Exception as exc:  # pragma: no cover - dependency declared
        raise ParserUnavailable("pypdf is required for .pdf files") from exc

    reader = PdfReader(str(path))
    sections: list[tuple[int | None, str | None, str | None, str]] = []
    for page_number, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            sections.append((page_number, None, f"page {page_number}", text))
    ocr_used = False
    if not sections and ocr:
        ocr_sections, ocr_warnings = _ocr_pdf_pages(path)
        sections.extend(ocr_sections)
        warnings.extend(ocr_warnings)
        ocr_used = bool(ocr_sections)
    if not sections:
        warnings.append(
            "No text extracted from PDF; scanned/OCR-only PDFs need `uv sync --extra ocr` and system OCR tools."
        )
    return _document_from_sections(
        stored_path, doc_id, "pdf-ocr" if ocr_used else "pdf", sections, warnings
    )


def _parse_image_ocr(
    path: Path, doc_id: str, stored_path: str, warnings: list[str]
) -> ParsedDocument:
    try:
        text = _ocr_image(path)
    except Exception as exc:  # noqa: BLE001 - convert dependency/tool failures to report warnings
        warnings.append(f"OCR failed for image: {exc}")
        text = ""
    if not text.strip():
        warnings.append("No OCR text extracted from image evidence.")
    return _document_from_sections(
        stored_path, doc_id, "image-ocr", [(None, None, path.stem, text)], warnings
    )


def _ocr_pdf_pages(
    path: Path,
) -> tuple[list[tuple[int | None, str | None, str | None, str]], list[str]]:
    try:
        from pdf2image import convert_from_path  # type: ignore[import-not-found]
    except Exception:
        return [], [
            "pdf2image is not installed; run `uv sync --extra ocr` to enable scanned PDF OCR."
        ]

    sections: list[tuple[int | None, str | None, str | None, str]] = []
    warnings: list[str] = []
    try:
        images = convert_from_path(str(path), poppler_path=_poppler_path())
        for page_number, image in enumerate(images, start=1):
            text = _ocr_image(image)
            if text.strip():
                sections.append((page_number, None, f"ocr page {page_number}", text))
    except Exception as exc:  # noqa: BLE001 - OCR tool availability varies by machine
        warnings.append(f"Scanned PDF OCR failed: {exc}")
    return sections, warnings


def _ocr_image(image_or_path: Any) -> str:
    try:
        import pytesseract  # type: ignore[import-not-found]
    except Exception as exc:  # pragma: no cover - optional dependency
        raise ParserUnavailable(
            "pytesseract is not installed; run `uv sync --extra ocr` and install Tesseract OCR."
        ) from exc
    command = _tesseract_cmd()
    if command:
        pytesseract.pytesseract.tesseract_cmd = command
    image_input = str(image_or_path) if isinstance(image_or_path, Path) else image_or_path
    return str(pytesseract.image_to_string(image_input))


def _tesseract_cmd() -> str | None:
    candidates = [
        os.environ.get("TESSERACT_CMD"),
        shutil.which("tesseract"),
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
    ]
    return next(
        (candidate for candidate in candidates if candidate and Path(candidate).exists()), None
    )


def _poppler_path() -> str | None:
    env_path = os.environ.get("POPPLER_PATH")
    if env_path and Path(env_path).exists():
        return env_path
    executable = shutil.which("pdftoppm")
    if executable:
        return str(Path(executable).parent)
    winget_root = Path.home() / "AppData/Local/Microsoft/WinGet/Packages"
    if winget_root.exists():
        for candidate in winget_root.glob("oschwartz10612.Poppler*/poppler-*/Library/bin"):
            if (candidate / "pdftoppm.exe").exists():
                return str(candidate)
    return None


def _document_from_sections(
    stored_path: str,
    doc_id: str,
    parser: str,
    sections: list[tuple[int | None, str | None, str | None, str]],
    warnings: list[str],
) -> ParsedDocument:
    chunks: list[ParsedChunk] = []
    for page, sheet, section, text in sections:
        for piece in chunk_text(text):
            chunks.append(
                ParsedChunk(
                    chunk_id=f"{doc_id}-chunk-{len(chunks) + 1:04d}",
                    doc_id=doc_id,
                    path=stored_path,
                    page=page,
                    sheet=sheet,
                    section=section,
                    text=piece,
                )
            )
    return ParsedDocument(path=stored_path, parser=parser, chunks=chunks, warnings=warnings)


def chunk_text(text: str, max_chars: int = MAX_CHARS, overlap: int = OVERLAP) -> list[str]:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    paragraphs = [
        paragraph.strip() for paragraph in re.split(r"\n{2,}", normalized) if paragraph.strip()
    ]
    if not paragraphs and normalized.strip():
        paragraphs = [normalized.strip()]

    chunks: list[str] = []
    current = ""
    for paragraph in paragraphs:
        if len(paragraph) > max_chars:
            if current:
                chunks.append(current.strip())
                current = ""
            chunks.extend(_split_long_text(paragraph, max_chars, overlap))
            continue
        candidate = f"{current}\n\n{paragraph}" if current else paragraph
        if len(candidate) <= max_chars:
            current = candidate
        else:
            chunks.append(current.strip())
            current = paragraph
    if current.strip():
        chunks.append(current.strip())
    return [normalize_ws(chunk) if "\n" not in chunk else chunk.strip() for chunk in chunks]


def _split_long_text(text: str, max_chars: int, overlap: int) -> list[str]:
    pieces: list[str] = []
    start = 0
    while start < len(text):
        end = min(len(text), start + max_chars)
        if end < len(text):
            boundary = max(text.rfind(". ", start, end), text.rfind("\n", start, end))
            if boundary > start + max_chars // 2:
                end = boundary + 1
        pieces.append(text[start:end].strip())
        if end >= len(text):
            break
        start = max(0, end - overlap)
    return pieces


def parsed_text(parsed: ParsedDocument) -> str:
    return "\n".join(chunk.text for chunk in parsed.chunks)


def as_jsonable(value: Any) -> Any:
    if hasattr(value, "model_dump"):
        return value.model_dump()
    return value
